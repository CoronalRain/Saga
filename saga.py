import requests
from lxml import html
import datetime
import inflect
from docx import Document
from docx.shared import Inches
from PyQt4 import QtGui, uic
from PyQt4.QtGui import QTableWidgetItem, QApplication, QColor, QIcon, QPushButton, QHBoxLayout, QAction
from PyQt4.QtCore import QDate, QTimer, QSize, Qt
import sys

# Start a new session.
session_requests = requests.session()

# Create some global URL variables.
login_url = "https://mindovermathtutoring.teachworks.com/accounts/login"
dashboard_url = "https://mindovermathtutoring.teachworks.com/dashboard"
participants_url = "https://mindovermathtutoring.teachworks.com/participants"

# Find the main UI elements.
if hasattr(sys, '_MEIPASS'):
    ui_login_path = os.path.join(sys._MEIPASS, "login_window.ui")
    ui_main_path = os.path.join(sys._MEIPASS, "main_window.ui")
else:
    ui_login_path = "login_window.ui"
    ui_main_path = "main_window.ui"

# Load the main UI elements.
Ui_LoginWindow, QLoginWindow = uic.loadUiType(ui_login_path)
Ui_MainWindow, QMainWindow = uic.loadUiType(ui_main_path)


class Lesson(object):
    def __init__(self, lesson_url, date, notes, topics):
        self.lesson_url = lesson_url
        self.date = date.strftime("%D")
        self.notes = notes
        self.topics = topics


class Student(object):
    def __init__(self, name, subject, tutor):
        self.name = name
        self.first_name = str.split(name)[0]
        self.last_name = str.split(name)[-1]
        self.subject = subject
        self.tutor = tutor
        self.lessons = []
        self.notes = []
        self.topics = []

    def __str__(self):
        st = "---------------\n"
        st += self.name + "\n"
        st += self.subject + "\n"
        st += "---------------\n"
        for i in range(len(self.lessons)):
            lesson = self.lessons[i]
            st += "Lesson " + lesson.code + "\n"
            st += lesson.date.isoformat() + "\n"
            st += lesson.notes + "\n"
        st += "Topics: "
        for topic in self.topics:
            st += topic + "; "
        st += "\n\n"

        return st

    def add_lesson(self, lesson):
        self.lessons.append(lesson)

    def get_topics(self):
        topics = []
        for lesson in self.lessons:
            topics.extend(lesson.topics)

        return topics


class RemoveButtonWidget(QtGui.QWidget):
    def __init__(self, parent, row):
        self.parent = parent
        self.row = row
        super(RemoveButtonWidget, self).__init__(parent)

        layout = QtGui.QHBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self.button = QtGui.QPushButton()
        self.button.setIcon(QIcon("gfx/x.png"))
        self.button.clicked.connect(self.removeRow)
        layout.addWidget(self.button)

        self.setLayout(layout)

    def setRow(self, row):
        self.row = row

    def removeRow(self):
        row = self.row
        self.parent.removeRow(row)
        self.updateRows()

    def updateRows(self):
        nrows = self.parent.rowCount()
        for i in range(nrows):
            self.parent.cellWidget(i, 0).setRow(i)
            self.parent.setItem(i, 1, QtGui.QTableWidgetItem(str(i+1)))


class LoginWindow(QLoginWindow, Ui_LoginWindow):
    def __init__(self, ):
        super(LoginWindow, self).__init__()
        self.setupUi(self)
        self.setWindowTitle("Mind Over Math Login")
        self.email_field.setText("tpk.mindovermath@gmail.com")
        self.password_field.setText("Redshift1!")

        self.login_button.clicked.connect(self.login_button_clicked)

    def login_button_clicked(self):
        # Get login information from user.
        email = self.email_field.text()
        password = self.password_field.text()

        # Open the login page.
        response = session_requests.get(login_url)

        # Grab hidden inputs.
        tree = html.fromstring(response.text)
        hidden_inputs = tree.xpath(r"//form//input[@type='hidden']")
        form = {x.attrib["name"]: x.attrib["value"] for x in hidden_inputs}

        # Enter username and password.
        form["user[email]"] = email
        form["user[password]"] = password

        # POST the login request.
        response = session_requests.post(login_url, data=form)
        if response.url == dashboard_url:
            self.error_label.setText("Login successful!")
            self.main_window = MainWindow()
            self.main_window.show()
            self.close()
        else:
            self.error_label.setText("Invalid email/password.")


class MainWindow(QMainWindow, Ui_MainWindow):
    def __init__(self, ):
        super(MainWindow, self).__init__()
        self.setupUi(self)
        self.setWindowTitle("Saga")

        # Set the column widths.
        self.student_table.setColumnWidth(0, 25)
        self.student_table.setColumnWidth(1, 25)
        self.student_table.setColumnWidth(2, 200)
        self.student_table.setColumnWidth(3, 200)
        self.student_table.setColumnWidth(4, 75)

        # Determine an appropriate date range for the progress reports.
        today = datetime.date.today()
        if today.day < 7:
            first = today.replace(day=1)
            last_month = first - datetime.timedelta(days=1)
            self.start_date_edit.setDate(QDate(last_month.year, last_month.month, 1))
            self.end_date_edit.setDate(QDate(last_month.year, last_month.month, last_month.day))
            self.names_ak_radio_button.toggle()
        elif today.day > 24:
            first = today.replace(day=1)
            next_month = first.replace(day=28) + datetime.timedelta(days=4)
            last_day = next_month - datetime.timedelta(days=next_month.day)
            self.start_date_edit.setDate(QDate(first.year, first.month, 1))
            self.end_date_edit.setDate(QDate(last_day.year, last_day.month, last_day.day))
            self.names_ak_radio_button.toggle()
        else:
            first = today.replace(day=1)
            last_month = first - datetime.timedelta(days=1)
            self.start_date_edit.setDate(QDate(last_month.year, last_month.month, 15))
            self.end_date_edit.setDate(QDate(today.year, today.month, 14))
            self.names_lz_radio_button.toggle()

        # Activate if something in the File menu was clicked.
        self.actionNew.triggered.connect(self.file_new)
        self.actionOpen.triggered.connect(self.file_open)
        self.actionSave.triggered.connect(self.file_save)

        # Activate if one of the GUI buttons was clicked.
        self.scrape_info_button.clicked.connect(self.scrape_info_button_clicked)
        self.generate_reports_button.clicked.connect(self.generate_reports_button_clicked)

    def file_new(self):
        # Remove all students from the table.
        self.student_table.setRowCount(0)

    def file_open(self):
        # Open a .saga file.
        print("Open triggered.")

    def file_save(self):
        # Save the table as a .saga file.
        print("Save triggered.")

    def scrape_info_button_clicked(self):
        # Reset the main GUI elements.
        self.progress_bar.setTextVisible(True)
        self.scrape_info_button.setEnabled(False)
        self.generate_reports_button.setEnabled(False)
        self.student_table.setRowCount(0)

        # Grab the user-selected ranges for dates and last names.
        start_date = self.start_date_edit.date().toPyDate()
        end_date = self.end_date_edit.date().toPyDate()
        valid_names = ""
        try:
            valid_names = self.valid_names_button_group.checkedButton().text()
        except AttributeError:
            self.error_label.setText("Please specify last names.")

        if valid_names != "":
            # Count the number of lessons and initiate the progress bar.
            num_lessons = self.count_lessons(start_date, end_date, valid_names)
            self.progress_bar.setMinimum(1)
            self.progress_bar.setMaximum(num_lessons)

            # Scrape the requested student info, and return a list of students.
            students = self.scrape_info(start_date, end_date, valid_names)

            for i in range(len(students)):
                # Grab each student.
                s = students[i]

                # Order the topics chronologically.
                topics = s.get_topics()
                topics.reverse()

                # Insert a new row in the table for each student.
                row_position = self.student_table.rowCount()
                self.student_table.insertRow(row_position)

                # Populate the row with relevant student info.
                self.student_table.setCellWidget(i, 0, RemoveButtonWidget(self.student_table, i))
                self.student_table.setItem(i, 1, QtGui.QTableWidgetItem(str(i+1)))
                self.student_table.setItem(i, 2, QtGui.QTableWidgetItem(s.name))
                self.student_table.setItem(i, 3, QtGui.QTableWidgetItem(s.subject))
                self.student_table.setItem(i, 4, QtGui.QTableWidgetItem(str(len(s.lessons))))
                self.student_table.setItem(i, 5, QtGui.QTableWidgetItem("\n".join(topics)))

                # Make sure the row fits inside the table.
                self.student_table.resizeRowToContents(row_position)

        # Reactivate the main GUI buttons.
        self.scrape_info_button.setEnabled(True)
        self.generate_reports_button.setEnabled(True)

    def count_lessons(self, start_date, end_date, valid_names):
        num_lessons = 0
        page = 0
        cont = True
        while cont:
            # Visit the next page of lessons.
            page += 1
            response = session_requests.get(participants_url + "?page=" + str(page))
            tree = html.fromstring(response.content)

            # If the last page is reached, break out of the loop.
            if tree.xpath("//div[@id='participants']/div/table/tbody/tr[1]/td[1]/text()")[0] == "No results found.":
                break
            else:
                # Scrape relevant information.
                dates = tree.xpath("//div[@id='participants']/div/table/tbody/tr/td[2]/text()")
                names = tree.xpath("//div[@id='participants']/div/table/tbody/tr/td[6]/text()")
                statuses = tree.xpath("//div[@id='participants']/div/table/tbody/tr/td[8]/text()")

                for j in range(len(dates)):
                    date = datetime.date(int(dates[j].split("/")[2]), int(dates[j].split("/")[0]),
                                         int(dates[j].split("/")[1]))
                    name = names[j]
                    status = statuses[j]

                    # Skip this entry if the date is outside of the acceptable range of values.
                    if date < start_date:
                        cont = False
                        break
                    if date > end_date:
                        continue

                    # Skip this entry if the last name is outside of the acceptable range of values.
                    letter = str.upper(name.split()[-1][0])
                    if (valid_names == "A-K" and letter > "K") or (valid_names == "L-Z" and letter < "L"):
                        continue

                    # Add one lesson to the running total.
                    if status == "Attended":
                        num_lessons += 1

        # Return the total number of lessons.
        return num_lessons

    def scrape_info(self, start_date, end_date, valid_names):
        students = []
        student_names = []
        num_lessons = 0
        page = 0
        cont = True
        while cont:
            page += 1

            # Visit the next page of lessons.
            response = session_requests.get(participants_url + "?page=" + str(page))
            tree = html.fromstring(response.content)

            # If the last page is reached, break out of the loop.
            if tree.xpath("//div[@id='participants']/div/table/tbody/tr[1]/td[1]/text()")[0] == "No results found.":
                break
            else:
                # Scrape relevant information.
                dates = tree.xpath("//div[@id='participants']/div/table/tbody/tr/td[2]/text()")
                lessons = tree.xpath("//div[@id='participants']/div/table/tbody/tr/td[4]/a/@href")
                subjects = tree.xpath("//div[@id='participants']/div/table/tbody/tr/td[4]/a/text()")
                tutors = tree.xpath("//div[@id='participants']/div/table/tbody/tr/td[5]/text()")
                names = tree.xpath("//div[@id='participants']/div/table/tbody/tr/td[6]/text()")
                statuses = tree.xpath("//div[@id='participants']/div/table/tbody/tr/td[8]/text()")

                for j in range(len(lessons)):
                    lesson_url = str.split(lessons[j], "/")[-1]
                    subject = str.split(subjects[j], "-")[-1].strip()
                    tutor = tutors[j]
                    name = names[j]
                    status = statuses[j]

                    # Skip this entry if the date is outside of the acceptable range of values.
                    date = datetime.date(int(dates[j].split("/")[2]), int(dates[j].split("/")[0]),
                                         int(dates[j].split("/")[1]))
                    if date < start_date:
                        cont = False
                        break
                    if date > end_date:
                        continue

                    # Skip this entry if the last name is outside of the acceptable range of values.
                    letter = str.upper(name.split()[-1][0])
                    if (valid_names == "A-K" and letter > "K") or (valid_names == "L-Z" and letter < "L"):
                        continue

                    # Either create a new student or add a lesson to an existing student.
                    if status == "Attended":
                        num_lessons += 1
                        student = None
                        if name not in student_names:
                            student = Student(name, subject, tutor)
                            student_names.append(name)
                            students.append(student)
                        else:
                            for s in students:
                                if s.name == name:
                                    student = s
                                    break

                        # Visit the lesson page and scrape the internal notes.
                        response = session_requests.get(participants_url + "/" + lesson_url)
                        tree = html.fromstring(response.content)
                        notes_array = tree.xpath("//div[@class='row participant-notes']/div[2]/span/text()")
                        notes = "\n".join(notes_array)

                        # If there are topics listed, parse them out.
                        notes_split = notes.split("Topics:")
                        topics = []
                        if len(notes_split) > 1:
                            topics = notes_split[-1].split(";")
                            topics = [topic.strip() for topic in topics]

                        # Add the new lesson for this student.
                        lesson = Lesson(lesson_url, date, notes, topics)
                        student.add_lesson(lesson)

                        # Update the progress bar.
                        self.progress_bar.setValue(num_lessons)
                        QApplication.processEvents()

        # Reset the progress bar.
        self.progress_bar.reset()
        self.progress_bar.setTextVisible(False)

        # Sort the list of students by last name.
        students.sort(key=lambda x: x.last_name)

        # Return the list of students.
        return students

    def generate_reports_button_clicked(self):
        # Create a new Word document.
        document = Document()

        # Initialize the progress bar.
        nrows = self.student_table.rowCount()
        self.progress_bar.setMinimum(1)
        self.progress_bar.setMaximum(nrows)
        self.progress_bar.setTextVisible(True)

        for row in range(nrows):
            # Grab relevant information from each row in the table.
            full_name = self.student_table.item(row, 2).text()
            subject = self.student_table.item(row, 3).text()
            num = self.student_table.item(row, 4).text()
            topics = self.student_table.item(row, 5).text().splitlines()

            # Parse some additional relevant information.
              ####################
              ### UPDATE THIS! ###
              ####################
            first_name = full_name.split(" ")[0]
            tutor = "Troy Kling"
            start_date = self.start_date_edit.date()
            end_date = self.end_date_edit.date()
            valid_names = ""
            try:
                valid_names = self.valid_names_button_group.checkedButton().text()
            except AttributeError:
                self.error_label.setText("Please specify last names.")

            # Write the intro paragraph for each student.
            p = document.add_paragraph("")
            p.add_run(full_name + " - " + subject).underline = True
            p.add_run("\n")
            p.add_run(
                "From " + start_date.toString("MMMM d") + " to " + end_date.toString("MMMM d") + ", " + first_name +
                " attended " + inflect.engine().number_to_words(num) + " tutoring sessions for " + subject + ". " +
                tutor.split()[0].strip() + " and " + first_name + " were able to cover the following topics:")

            # Write the bullet list of topics for each student.
            for topic in topics:
                document.add_paragraph(topic, style="List Bullet").paragraph_format.left_indent = Inches(.5)

            # Write a filler paragraph at the end (to be manually completed by the tutor).
            document.add_paragraph(
                "Paragraph detailing: (1) improvements, breakthroughs; (2) struggles, solutions; (3) test grades, " +
                "positive/negative, plan of action; (4) concerns about student; (5) goals for future sessions. " +
                "Please let us know if you have any questions about " + first_name + "'s progress.\n")

            # Update the progress bar.
            self.progress_bar.setValue(row)
            QApplication.processEvents()

            # Add a page break every couple students.
            if row % 2 == 1:
                document.add_page_break()

        # Save the progress reports as a Word document.
        document.save(start_date.toString("MMMM yyyy") + " (" + valid_names + ")" + " Progress Reports " +
                      tutor.split()[0].strip()[0] + " " + tutor.split()[-1].strip() + ".docx")

        # Reset the progress bar.
        self.progress_bar.reset()
        self.progress_bar.setTextVisible(False)

    def closeEvent(self, event):
        # Close the application nicely.
        self.deleteLater()


if __name__ == "__main__":
    # Start the application.
    app = QtGui.QApplication(sys.argv)

    # Set icons of various sizes.
    app_icon = QtGui.QIcon()
    app_icon.addFile("gfx/saga16.png", QSize(16, 16))
    app_icon.addFile("gfx/saga24.png", QSize(24, 24))
    app_icon.addFile("gfx/saga32.png", QSize(32, 32))
    app_icon.addFile("gfx/saga64.png", QSize(64, 64))
    app_icon.addFile("gfx/saga256.png", QSize(256, 256))
    app.setWindowIcon(app_icon)

    # Open the login window.
    login_window = LoginWindow()
    login_window.show()

    # Terminate when finished.
    sys.exit(app.exec_())
